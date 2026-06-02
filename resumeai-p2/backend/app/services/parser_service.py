"""
parser_service.py
Converts uploaded PDF or DOCX files into a structured JSON schema
that every AI analyzer in the pipeline reads from.

Output schema:
{
  "contact": { name, email, phone, linkedin, location },
  "summary": "...",
  "experience": [
    { company, title, start, end, bullets: ["..."] }
  ],
  "education": [
    { institution, degree, field, start, end, gpa }
  ],
  "skills": ["Python", "React", ...],
  "certifications": ["..."],
  "projects": [{ name, description, bullets }],
  "raw_text": "..."
}
"""
import re
import io
from pathlib import Path
from typing import BinaryIO


# ── PDF Parsing ───────────────────────────────────────────────────────────────

def parse_pdf(file: BinaryIO) -> dict:
    """Extract and structure text from a PDF file."""
    import pdfplumber

    full_text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if text:
                full_text += text + "\n"

    # If pdfplumber returns nothing, file is likely scanned — try OCR
    if not full_text.strip():
        full_text = _ocr_pdf(file)

    return _structure_text(full_text)


def _ocr_pdf(file: BinaryIO) -> str:
    """Fallback OCR for image-based or scanned PDFs."""
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        images = convert_from_bytes(file.read())
        return "\n".join(pytesseract.image_to_string(img) for img in images)
    except Exception:
        return ""


# ── DOCX Parsing ──────────────────────────────────────────────────────────────

def parse_docx(file: BinaryIO) -> dict:
    """Extract and structure text from a DOCX file."""
    from docx import Document

    doc = Document(file)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Also extract from tables (skills tables, etc.)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in lines:
                    lines.append(text)

    full_text = "\n".join(lines)
    return _structure_text(full_text)


# ── Text Structuring ──────────────────────────────────────────────────────────

def _structure_text(raw_text: str) -> dict:
    """
    Heuristic section detection.
    For production, the AI pipeline re-parses this with Claude for higher accuracy,
    but this gives us a clean base structure without burning tokens.
    """
    lines = [l.strip() for l in raw_text.split("\n") if l.strip()]

    sections = _split_into_sections(lines)

    return {
        "contact":        _parse_contact(sections.get("header", [])),
        "summary":        _parse_summary(sections.get("summary", [])),
        "experience":     _parse_experience(sections.get("experience", [])),
        "education":      _parse_education(sections.get("education", [])),
        "skills":         _parse_skills(sections.get("skills", [])),
        "certifications": _parse_certifications(sections.get("certifications", [])),
        "projects":       _parse_projects(sections.get("projects", [])),
        "raw_text":       raw_text,
    }


# Section header keywords (case-insensitive)
SECTION_PATTERNS = {
    "summary":        r"^(summary|profile|objective|about|professional\s+summary|career\s+summary)",
    "experience":     r"^(experience|work\s+experience|employment|work\s+history|professional\s+experience|career)",
    "education":      r"^(education|academic|qualifications|degrees?)",
    "skills":         r"^(skills?|technical\s+skills?|core\s+competencies|competencies|technologies|tools)",
    "certifications": r"^(certifications?|licenses?|credentials?|accreditations?)",
    "projects":       r"^(projects?|personal\s+projects?|portfolio|open\s+source)",
    "awards":         r"^(awards?|honors?|achievements?|recognition)",
}


def _split_into_sections(lines: list[str]) -> dict[str, list[str]]:
    """Split resume lines into labelled sections."""
    sections: dict[str, list[str]] = {"header": []}
    current_section = "header"
    header_done = False

    for line in lines:
        # Check if this line is a section header
        matched_section = None
        for section_name, pattern in SECTION_PATTERNS.items():
            if re.match(pattern, line, re.IGNORECASE) and len(line) < 60:
                matched_section = section_name
                break

        if matched_section:
            current_section = matched_section
            header_done = True
            if current_section not in sections:
                sections[current_section] = []
        else:
            if not header_done and len(sections["header"]) < 8:
                sections["header"].append(line)
            else:
                if current_section not in sections:
                    sections[current_section] = []
                sections[current_section].append(line)

    return sections


def _parse_contact(header_lines: list[str]) -> dict:
    contact = {
        "name": "",
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
        "location": "",
    }

    full_header = " ".join(header_lines)

    # Email
    email_match = re.search(r"[\w.+-]+@[\w-]+\.[a-z]{2,}", full_header, re.IGNORECASE)
    if email_match:
        contact["email"] = email_match.group(0)

    # Phone
    phone_match = re.search(
        r"(\+?\d[\d\s\-().]{7,}\d)", full_header
    )
    if phone_match:
        contact["phone"] = phone_match.group(0).strip()

    # LinkedIn
    linkedin_match = re.search(r"linkedin\.com/in/([\w-]+)", full_header, re.IGNORECASE)
    if linkedin_match:
        contact["linkedin"] = f"linkedin.com/in/{linkedin_match.group(1)}"

    # GitHub
    github_match = re.search(r"github\.com/([\w-]+)", full_header, re.IGNORECASE)
    if github_match:
        contact["github"] = f"github.com/{github_match.group(1)}"

    # Name = first non-empty line that is not an email/phone/URL and is title-like
    for line in header_lines:
        line = line.strip()
        if (
            line
            and not re.search(r"[@|/\\]", line)
            and not re.search(r"\d{5,}", line)
            and len(line.split()) <= 6
            and len(line) > 2
        ):
            contact["name"] = line
            break

    # Location: look for City, State / City, Country patterns
    loc_match = re.search(r"([A-Za-z\s]+,\s*[A-Z]{2,})", full_header)
    if loc_match:
        contact["location"] = loc_match.group(0).strip()

    return contact


def _parse_summary(lines: list[str]) -> str:
    return " ".join(lines).strip()


def _parse_experience(lines: list[str]) -> list[dict]:
    experiences = []
    current: dict | None = None

    for line in lines:
        # Date pattern: detect job entry starts
        date_pattern = re.search(
            r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|january|february|"
            r"march|april|june|july|august|september|october|november|december"
            r"|\d{4})",
            line, re.IGNORECASE
        )

        # Bullet points
        is_bullet = re.match(r"^[•\-\*▪→]\s+", line) or re.match(r"^\d+\.\s+", line)

        if date_pattern and not is_bullet and len(line) < 150:
            # Likely a job title / company line
            if current:
                experiences.append(current)
            current = {
                "company":  "",
                "title":    line,
                "start":    "",
                "end":      "",
                "bullets":  [],
                "raw":      line,
            }
            # Try to extract dates
            dates = re.findall(
                r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|january|"
                r"february|march|april|june|july|august|september|october|november|"
                r"december|\d{4}|present|current)",
                line, re.IGNORECASE
            )
            if len(dates) >= 2:
                current["start"] = dates[0]
                current["end"]   = dates[-1]
            elif len(dates) == 1:
                current["start"] = dates[0]

        elif is_bullet and current:
            bullet_text = re.sub(r"^[•\-\*▪→]\s+|^\d+\.\s+", "", line).strip()
            current["bullets"].append(bullet_text)

        elif current and not is_bullet and line:
            # Could be company name or continuation
            if not current["company"] and len(line) < 100:
                current["company"] = line

    if current:
        experiences.append(current)

    return experiences


def _parse_education(lines: list[str]) -> list[dict]:
    educations = []
    current: dict | None = None

    degree_keywords = r"(bachelor|master|phd|ph\.d|b\.s|b\.a|m\.s|m\.a|mba|associate|diploma|degree|b\.tech|m\.tech)"

    for line in lines:
        if re.search(degree_keywords, line, re.IGNORECASE) or re.search(r"\d{4}", line):
            if current:
                educations.append(current)
            current = {
                "institution": "",
                "degree":      line,
                "field":       "",
                "start":       "",
                "end":         "",
                "gpa":         "",
            }
            # Extract GPA
            gpa_match = re.search(r"gpa[:\s]*([\d.]+)", line, re.IGNORECASE)
            if gpa_match:
                current["gpa"] = gpa_match.group(1)

            # Extract years
            years = re.findall(r"\d{4}", line)
            if len(years) >= 2:
                current["start"] = years[0]
                current["end"]   = years[1]
            elif len(years) == 1:
                current["end"] = years[0]

        elif current and not current["institution"]:
            current["institution"] = line

    if current:
        educations.append(current)

    return educations


def _parse_skills(lines: list[str]) -> list[str]:
    skills = []
    for line in lines:
        # Split on common delimiters
        parts = re.split(r"[,|•·/\t]+", line)
        for part in parts:
            skill = part.strip()
            # Filter out section headers and very long strings
            if skill and 1 < len(skill) < 50 and not re.match(r"skills?", skill, re.IGNORECASE):
                skills.append(skill)
    return list(dict.fromkeys(skills))  # deduplicate preserving order


def _parse_certifications(lines: list[str]) -> list[str]:
    return [line.strip() for line in lines if line.strip()]


def _parse_projects(lines: list[str]) -> list[dict]:
    projects = []
    current: dict | None = None

    for line in lines:
        is_bullet = re.match(r"^[•\-\*▪→]\s+", line)
        if not is_bullet and line and len(line) < 120:
            if current:
                projects.append(current)
            current = {"name": line, "description": "", "bullets": []}
        elif is_bullet and current:
            bullet_text = re.sub(r"^[•\-\*▪→]\s+", "", line).strip()
            current["bullets"].append(bullet_text)
        elif current and not is_bullet:
            current["description"] += " " + line

    if current:
        projects.append(current)

    return projects
